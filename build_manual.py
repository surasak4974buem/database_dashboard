#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
สร้าง Word (.docx) และ PDF พร้อมกัน
ฟอนต์ TH Sarabun New ขนาดตามมาตรฐาน HA / ISO 13485
กระดาษ A4 จัดหน้าสวยงาม มีรูปประกอบ
"""

import io, os, sys, urllib.request
from pathlib import Path

# ─── paths ────────────────────────────────────────────────────────────────────
BASE   = Path("/home/user/database_dashboard")
FONTS  = Path("/tmp/fonts")
FONTS.mkdir(parents=True, exist_ok=True)
IMGS   = BASE / "assets"
IMGS.mkdir(parents=True, exist_ok=True)

WORD_OUT = BASE / "คู่มือตรวจเช็คงานซ่อมเครื่องมือแพทย์.docx"
PDF_OUT  = BASE / "คู่มือตรวจเช็คงานซ่อมเครื่องมือแพทย์.pdf"

# ─── download Sarabun fonts ────────────────────────────────────────────────────
FONT_URLS = {
    "Sarabun-Regular": "https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Regular.ttf",
    "Sarabun-Bold":    "https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Bold.ttf",
    "Sarabun-Italic":  "https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Italic.ttf",
}
for name, url in FONT_URLS.items():
    path = FONTS / f"{name}.ttf"
    if not path.exists():
        print(f"  downloading {name}...")
        urllib.request.urlretrieve(url, path)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — สร้างรูป icon ด้วย Pillow
# ══════════════════════════════════════════════════════════════════════════════
from PIL import Image, ImageDraw, ImageFont
import textwrap

DARK_BLUE  = (0,   51, 102)
MED_BLUE   = (0,   90, 156)
LIGHT_BLUE = (214, 228, 240)
RED        = (192,  57,  43)
ORANGE     = (230, 126,  34)
GREEN      = (30,  132,  73)
WHITE      = (255, 255, 255)
YELLOW     = (255, 243, 205)

def icon(color_bg, text, size=128):
    img = Image.new("RGBA", (size, size), (0,0,0,0))
    d   = ImageDraw.Draw(img)
    r   = size // 2
    d.ellipse([0, 0, size-1, size-1], fill=color_bg)
    try:
        fnt = ImageFont.truetype(str(FONTS/"Sarabun-Bold.ttf"), int(size*0.42))
    except Exception:
        fnt = ImageFont.load_default()
    bbox = d.textbbox((0,0), text, font=fnt)
    tx = (size - (bbox[2]-bbox[0])) // 2
    ty = (size - (bbox[3]-bbox[1])) // 2
    d.text((tx, ty), text, fill=WHITE, font=fnt)
    return img

def save_png(img, name):
    p = IMGS / name
    img.save(p, "PNG")
    return p

# รูป banner ส่วนหัว
def make_banner():
    W, H = 1754, 200   # A4 width @150dpi approx
    img = Image.new("RGB", (W, H), DARK_BLUE)
    d   = ImageDraw.Draw(img)
    # gradient strip
    for i in range(W):
        ratio = i / W
        r = int(DARK_BLUE[0] + (MED_BLUE[0]-DARK_BLUE[0])*ratio)
        g = int(DARK_BLUE[1] + (MED_BLUE[1]-DARK_BLUE[1])*ratio)
        b = int(DARK_BLUE[2] + (MED_BLUE[2]-DARK_BLUE[2])*ratio)
        d.line([(i,0),(i,H)], fill=(r,g,b))
    # accent bar bottom
    d.rectangle([0, H-12, W, H], fill=ORANGE)
    # text
    try:
        f_big  = ImageFont.truetype(str(FONTS/"Sarabun-Bold.ttf"),    72)
        f_small= ImageFont.truetype(str(FONTS/"Sarabun-Regular.ttf"), 40)
    except Exception:
        f_big = f_small = ImageFont.load_default()
    line1 = "คู่มือปฏิบัติงาน — การตรวจเช็คงานซ่อมเครื่องมือแพทย์"
    line2 = "หน่วยวิศวกรรมการแพทย์  |  08:30 – 16:30 น.  |  จันทร์ – ศุกร์ (วันราชการ)"
    b1 = d.textbbox((0,0), line1, font=f_big)
    b2 = d.textbbox((0,0), line2, font=f_small)
    d.text(((W-(b1[2]-b1[0]))//2, 28),  line1, font=f_big,   fill=WHITE)
    d.text(((W-(b2[2]-b2[0]))//2, 118), line2, font=f_small, fill=(214,228,240))
    return save_png(img, "banner.png")

# รูปไอคอนแต่ละส่วน
def make_section_icons():
    icons_def = [
        ("icon_d1.png",    RED,      "D1"),
        ("icon_d2.png",    ORANGE,   "D2"),
        ("icon_d3.png",    GREEN,    "D3"),
        ("icon_pm.png",    (100,100,100), "PM"),
        ("icon_step.png",  MED_BLUE, "✓"),
        ("icon_risk.png",  RED,      "!"),
        ("icon_form.png",  DARK_BLUE,"📋"),
        ("icon_esc.png",   ORANGE,   "↑"),
    ]
    paths = {}
    for fname, color, txt in icons_def:
        paths[fname] = save_png(icon(color, txt, 96), fname)
    return paths

# รูปตาราง Timeline ช่วงเวลาทำงาน
def make_timeline():
    W, H = 1600, 160
    img = Image.new("RGB", (W, H), WHITE)
    d   = ImageDraw.Draw(img)
    slots = [
        ("08:30", "09:00", "Morning Round\nตรวจงาน",            MED_BLUE),
        ("09:00", "12:00", "ดำเนินการซ่อม\nงาน D1/D2",         RED),
        ("12:00", "13:00", "พักกลางวัน",                        (180,180,180)),
        ("13:00", "15:30", "ดำเนินการซ่อม\nงาน D2/D3/PM",      GREEN),
        ("15:30", "16:00", "สรุปงาน\nอัปเดตระบบ",              ORANGE),
        ("16:00", "16:30", "ส่งมอบงาน\nรายงานหัวหน้า",         DARK_BLUE),
    ]
    total_min = (16*60+30) - (8*60+30)
    margin = 40
    bar_w  = W - 2*margin
    bar_y1, bar_y2 = 70, 120
    try:
        fnt_s = ImageFont.truetype(str(FONTS/"Sarabun-Regular.ttf"), 22)
        fnt_t = ImageFont.truetype(str(FONTS/"Sarabun-Bold.ttf"),    24)
    except Exception:
        fnt_s = fnt_t = ImageFont.load_default()

    def to_min(t):
        h,m = map(int, t.split(":"))
        return h*60+m

    start_m = to_min("08:30")
    for (t1, t2, label, col) in slots:
        x1 = margin + int((to_min(t1)-start_m) / total_min * bar_w)
        x2 = margin + int((to_min(t2)-start_m) / total_min * bar_w)
        d.rectangle([x1, bar_y1, x2-2, bar_y2], fill=col)
        cx = (x1+x2)//2
        lines = label.split("\n")
        for i, ln in enumerate(lines):
            bb = d.textbbox((0,0), ln, font=fnt_s)
            tx = cx - (bb[2]-bb[0])//2
            d.text((tx, bar_y1-28+i*22), ln, font=fnt_s, fill=col if i==0 else (80,80,80))
        d.text((x1, bar_y2+4), t1, font=fnt_s, fill=(80,80,80))

    bb = d.textbbox((0,0), "16:30", font=fnt_s)
    d.text((W-margin-(bb[2]-bb[0]), bar_y2+4), "16:30", font=fnt_s, fill=(80,80,80))

    title = "ตารางการปฏิบัติงานประจำวัน (08:30 – 16:30 น.)"
    bb = d.textbbox((0,0), title, font=fnt_t)
    d.text(((W-(bb[2]-bb[0]))//2, 5), title, font=fnt_t, fill=DARK_BLUE)
    return save_png(img, "timeline.png")

print("สร้างรูปภาพประกอบ...")
banner_path  = make_banner()
icon_paths   = make_section_icons()
timeline_path= make_timeline()
print("  ✓ รูปภาพเสร็จ")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — สร้าง Word Document
# มาตรฐาน HA/ISO: ฟอนต์ TH Sarabun New, เนื้อหา 16pt, หัวข้อ 18-20pt
# ══════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

doc = Document()

# ── Page setup A4 ──────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.0)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.0)

# ── helper: set paragraph font ─────────────────────────────────────────────
FONT_THAI = "TH Sarabun New"

def para_fmt(para, size_pt=16, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT,
             color=None, space_before=0, space_after=6,
             line_spacing=1.5):
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    for run in para.runs:
        run.font.name = FONT_THAI
        run.font.size = Pt(size_pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        # force Thai complex script font
        rpr = run._r.get_or_add_rPr()
        cs = OxmlElement('w:cs')
        cs.set(qn('w:val'), FONT_THAI)
        rPrCs = rpr.find(qn('w:cs'))
        if rPrCs is not None:
            rpr.remove(rPrCs)
        rpr.append(cs)
    return para

def add_heading(doc, text, level=1, page_break=False):
    if page_break:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(docx.oxml.OxmlElement('w:lastRenderedPageBreak'))
        p.paragraph_format.space_after = Pt(0)
    sizes = {1: (20, DARK_BLUE), 2: (18, MED_BLUE), 3: (16, DARK_BLUE)}
    sz, col = sizes.get(level, (16, DARK_BLUE))
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = FONT_THAI
    run.font.size  = Pt(sz)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*col)
    rpr = run._r.get_or_add_rPr()
    cs = OxmlElement('w:cs'); cs.set(qn('w:val'), FONT_THAI)
    rPrCs = rpr.find(qn('w:cs'))
    if rPrCs is not None: rpr.remove(rPrCs)
    rpr.append(cs)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10 if level==1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    # bottom border for H1
    if level == 1:
        ppr  = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '003366')
        pbdr.append(bot)
        ppr.append(pbdr)
    return p

def add_body(doc, text, indent=False, bullet=False):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
        prefix = "  "
    else:
        prefix = ""
    run = p.add_run(prefix + text)
    run.font.name = FONT_THAI
    run.font.size = Pt(16)
    rpr = run._r.get_or_add_rPr()
    cs = OxmlElement('w:cs'); cs.set(qn('w:val'), FONT_THAI)
    rPrCs = rpr.find(qn('w:cs'))
    if rPrCs is not None: rpr.remove(rPrCs)
    rpr.append(cs)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name  = FONT_THAI
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._r.get_or_add_rPr()
    cs = OxmlElement('w:cs'); cs.set(qn('w:val'), FONT_THAI)
    rPrCs = rpr.find(qn('w:cs'))
    if rPrCs is not None: rpr.remove(rPrCs)
    rpr.append(cs)

def add_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BDBDBD')
        tblBorders.append(el)
    tblPr.append(tblBorders)

# ══════════════════════════════════════════════════════════════════════════════
# BANNER IMAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_picture(str(banner_path), width=Cm(16.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # spacer

# ── Meta info table ─────────────────────────────────────────────────────────
meta_t = doc.add_table(rows=3, cols=4)
meta_t.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_t.style = 'Table Grid'
add_table_borders(meta_t)
meta_rows = [
    ("ฉบับที่",       "001",                                        "วันที่บังคับใช้", "มิถุนายน 2569"),
    ("แผนก",          "วิศวกรรมการแพทย์ / ซ่อมบำรุงเครื่องมือแพทย์", "เวลาปฏิบัติงาน", "08:30 – 16:30 น."),
    ("วันทำการ",      "จันทร์ – ศุกร์ (วันราชการ)",                  "ผู้รับผิดชอบ",   "ช่างเครื่องมือแพทย์ /\nวิศวกรชีวการแพทย์"),
]
col_w = [Cm(2.8), Cm(6.2), Cm(3.6), Cm(3.9)]
for i, row_data in enumerate(meta_rows):
    cells = meta_t.rows[i].cells
    for j, (txt, w) in enumerate(zip(row_data, col_w)):
        is_label = j % 2 == 0
        set_cell_bg(cells[j], "D6E4F0" if is_label else "FFFFFF")
        cell_text(cells[j], txt, size=14, bold=is_label,
                  align=WD_ALIGN_PARAGRAPH.CENTER if is_label else WD_ALIGN_PARAGRAPH.LEFT,
                  color=DARK_BLUE if is_label else None)
        cells[j].width = col_w[j]

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — วัตถุประสงค์
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. วัตถุประสงค์")
add_body(doc, "เพื่อคัดกรองและจัดลำดับความสำคัญของงานซ่อมเครื่องมือแพทย์ก่อนเริ่มปฏิบัติงานทุกวัน ให้งานที่เร่งด่วนและส่งผลต่อความปลอดภัยของผู้ป่วยได้รับการแก้ไขก่อนเป็นลำดับแรก ตามมาตรฐาน HA และ ISO 13485")

# ══════════════════════════════════════════════════════════════════════════════
# TIMELINE IMAGE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. ตารางปฏิบัติงานประจำวัน")
doc.add_picture(str(timeline_path), width=Cm(16.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ระดับความเร่งด่วน
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. เกณฑ์จำแนกระดับความเร่งด่วน")

urg_t = doc.add_table(rows=5, cols=5)
urg_t.alignment = WD_TABLE_ALIGNMENT.CENTER
urg_t.style = 'Table Grid'
add_table_borders(urg_t)
urg_cols = [Cm(2.0), Cm(1.5), Cm(6.5), Cm(3.5), Cm(3.0)]
urg_header = ["ระดับ","รหัส","ความหมาย","ระยะเวลา","ช่วงเวลา"]
urg_data = [
    ("วิกฤต",    "D1", "เครื่องมือเกี่ยวข้องกับชีวิตผู้ป่วยโดยตรง หรือหน่วย ICU/ER/OR", "ภายใน 1 ชั่วโมง",    "09:00–10:00"),
    ("เร่งด่วน", "D2", "กระทบการรักษา/วินิจฉัย แต่มีเครื่องสำรอง",                      "ภายใน 4 ชั่วโมง",    "ก่อน 12:00"),
    ("ปกติ",     "D3", "ไม่กระทบการรักษาในทันที",                                         "ภายใน 24–48 ชม.",    "ตามลำดับ"),
    ("แผน PM",   "PM", "งาน Preventive Maintenance ตามกำหนดการ",                          "ตามแผนที่กำหนด",    "ตามแผน"),
]
urg_bg = ["003366","C0392B","FEF9E7","EAFAF1","F2F3F4"]
urg_tc = [WHITE,   WHITE,    None,    None,    None  ]

for j, (h, w) in enumerate(zip(urg_header, urg_cols)):
    set_cell_bg(urg_t.rows[0].cells[j], "003366")
    cell_text(urg_t.rows[0].cells[j], h, size=14, bold=True, color=WHITE)

row_bg = ["C0392B","FEF5E7","EAFAF1","F2F3F4"]
for i, (row_d, bg) in enumerate(zip(urg_data, row_bg)):
    for j, txt in enumerate(row_d):
        set_cell_bg(urg_t.rows[i+1].cells[j], bg)
        cell_text(urg_t.rows[i+1].cells[j], txt, size=14,
                  bold=(j<=1 and i==0),
                  color=WHITE if i==0 else None,
                  align=WD_ALIGN_PARAGRAPH.CENTER if j!=2 else WD_ALIGN_PARAGRAPH.LEFT)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — ขั้นตอน Morning Round
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. ขั้นตอนการตรวจเช็คตอนเช้า — Morning Round (08:30 – 09:00 น.)")

steps = [
    ("ขั้นที่ 1", "ตรวจสอบใบแจ้งซ่อมค้างจากวันก่อน",
     ["เปิดระบบแจ้งซ่อม (CMMS / ระบบโรงพยาบาล)",
      "ตรวจสอบงานค้างที่ยังไม่เสร็จ และงานที่รอชิ้นส่วนอะไหล่",
      "บันทึกสถานะล่าสุดของงานค้างทุกรายการ"]),
    ("ขั้นที่ 2", "รับและคัดกรองใบแจ้งซ่อมใหม่",
     ["ตรวจสอบระบบออนไลน์ / สมุดรับแจ้งซ่อม / LINE / E-mail",
      "คัดกรองความเร่งด่วนด้วย 4 คำถาม (ดูหัวข้อ 5)",
      "บันทึกและกำหนดระดับ D1–D3 หรือ PM"]),
    ("ขั้นที่ 3", "จัดทำ Daily Work List จัดเรียง D1 → D2 → D3 → PM",
     ["ระบุชื่อเครื่องมือ หน่วยงาน ผู้แจ้งซ่อม และระดับความเร่งด่วน",
      "ประเมินเวลาที่ใช้และจัดสรรช่างให้เหมาะสม"]),
    ("ขั้นที่ 4", "มอบหมายงานและเตรียมทรัพยากร",
     ["มอบหมายช่างตามความเชี่ยวชาญและประเภทเครื่องมือ",
      "ตรวจสอบความพร้อมของเครื่องมือช่าง และอะไหล่",
      "แจ้งหน่วยงานผู้ใช้ทราบกำหนดการเข้าซ่อม"]),
    ("ขั้นที่ 5", "สรุปและรายงาน",
     ["บันทึกงานลงระบบ CMMS หรือแบบฟอร์ม ME-F-01",
      "รายงานงาน D1 ต่อหัวหน้าหน่วยทันที",
      "อัปเดตสถานะงานในระบบทุก 2 ชั่วโมง"]),
]

for step_no, title, bullets in steps:
    # colored step header
    step_t = doc.add_table(rows=1, cols=2)
    step_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    step_t.style = 'Table Grid'
    set_cell_bg(step_t.rows[0].cells[0], "005A9C")
    set_cell_bg(step_t.rows[0].cells[1], "D6E4F0")
    cell_text(step_t.rows[0].cells[0], step_no, size=14, bold=True, color=WHITE)
    cell_text(step_t.rows[0].cells[1], title, size=15, bold=True, color=DARK_BLUE,
              align=WD_ALIGN_PARAGRAPH.LEFT)
    step_t.rows[0].cells[0].width = Cm(2.0)
    step_t.rows[0].cells[1].width = Cm(14.5)
    for b in bullets:
        add_body(doc, f"✓  {b}", indent=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — คำถามคัดกรอง
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. คำถามคัดกรองความเร่งด่วน (ถามผู้แจ้งซ่อม)")

q_t = doc.add_table(rows=5, cols=3)
q_t.alignment = WD_TABLE_ALIGNMENT.CENTER
q_t.style = 'Table Grid'
add_table_borders(q_t)
set_cell_bg(q_t.rows[0].cells[0], "003366")
set_cell_bg(q_t.rows[0].cells[1], "003366")
set_cell_bg(q_t.rows[0].cells[2], "003366")
cell_text(q_t.rows[0].cells[0], "ข้อ",          size=14, bold=True, color=WHITE)
cell_text(q_t.rows[0].cells[1], "คำถาม",         size=14, bold=True, color=WHITE)
cell_text(q_t.rows[0].cells[2], "ตอบ ใช่ → ระดับ", size=14, bold=True, color=WHITE)
q_data = [
    ("1", "เครื่องมือนี้ใช้กับผู้ป่วยโดยตรงขณะนี้หรือไม่?",          "D1"),
    ("2", "อยู่ในหน่วย ICU / ER / OR / NICU หรือไม่?",               "D1"),
    ("3", "มีผู้ป่วยรอใช้งานและไม่มีเครื่องสำรองหรือไม่?",            "D1"),
    ("4", "กระทบการวินิจฉัย/รักษา แต่มีเครื่องสำรองหรือไม่?",         "D2"),
]
for i, (no, q, lv) in enumerate(q_data):
    bg = "FDEDEC" if lv=="D1" else "FEF5E7"
    for j in range(3):
        set_cell_bg(q_t.rows[i+1].cells[j], bg)
    cell_text(q_t.rows[i+1].cells[0], no, size=14)
    cell_text(q_t.rows[i+1].cells[1], q,  size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(q_t.rows[i+1].cells[2], lv, size=14, bold=True,
              color=RED if lv=="D1" else ORANGE)
q_t.rows[0].cells[0].width = Cm(1.0)
q_t.rows[0].cells[1].width = Cm(12.0)
q_t.rows[0].cells[2].width = Cm(3.5)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — เครื่องมือกลุ่มเสี่ยงสูง
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. เครื่องมือแพทย์กลุ่มเสี่ยงสูง (แก้ไขก่อนเสมอ)")

hi = [
    ["เครื่องช่วยหายใจ (Ventilator)",           "เครื่องกระตุกหัวใจ (Defibrillator / AED)"],
    ["เครื่องดมยาสลบ (Anesthesia Machine)",       "เครื่องให้ยาอัตโนมัติ (Infusion / Syringe Pump)"],
    ["เครื่องฟอกไต (Dialysis Machine)",           "เครื่องตรวจติดตามผู้ป่วย (Patient Monitor)"],
    ["เครื่อง X-ray / CT Scan (กรณีด่วน)",       "เครื่องให้ออกซิเจน (O2 Concentrator / Flow Meter)"],
]
hi_t = doc.add_table(rows=len(hi), cols=2)
hi_t.alignment = WD_TABLE_ALIGNMENT.CENTER
hi_t.style = 'Table Grid'
add_table_borders(hi_t)
for i, row_data in enumerate(hi):
    for j, txt in enumerate(row_data):
        set_cell_bg(hi_t.rows[i].cells[j], "FDEDEC")
        cell_text(hi_t.rows[i].cells[j], f"⚕  {txt}", size=14,
                  align=WD_ALIGN_PARAGRAPH.LEFT, color=(100,0,0))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Escalation
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. เกณฑ์การยกระดับความเร่งด่วน (Escalation)")

esc_data = [
    ("งาน D2 ที่ยังไม่เสร็จใน 4 ชั่วโมง",    "ยกระดับเป็น D1 และรายงานหัวหน้าทันที"),
    ("ไม่มีเครื่องสำรอง + งาน D2",            "ยกระดับเป็น D1 ทันที"),
    ("ต้องส่งซ่อมภายนอก",                     "แจ้งหัวหน้าและติดต่อบริษัทภายในวันนั้น"),
    ("งาน D3 ค้างเกิน 3 วันทำการ",            "ยกระดับเป็น D2 และรายงานหัวหน้า"),
]
esc_t = doc.add_table(rows=len(esc_data)+1, cols=2)
esc_t.alignment = WD_TABLE_ALIGNMENT.CENTER
esc_t.style = 'Table Grid'
add_table_borders(esc_t)
set_cell_bg(esc_t.rows[0].cells[0], "003366")
set_cell_bg(esc_t.rows[0].cells[1], "003366")
cell_text(esc_t.rows[0].cells[0], "สถานการณ์",   size=14, bold=True, color=WHITE)
cell_text(esc_t.rows[0].cells[1], "การดำเนินการ", size=14, bold=True, color=WHITE)
for i, (sit, act) in enumerate(esc_data):
    bg = "FEF9E7" if i%2==0 else "FFFFFF"
    set_cell_bg(esc_t.rows[i+1].cells[0], bg)
    set_cell_bg(esc_t.rows[i+1].cells[1], bg)
    cell_text(esc_t.rows[i+1].cells[0], sit, size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(esc_t.rows[i+1].cells[1], act, size=14, bold=True,
              align=WD_ALIGN_PARAGRAPH.LEFT, color=RED)
esc_t.rows[0].cells[0].width = Cm(8.0)
esc_t.rows[0].cells[1].width = Cm(8.5)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — แบบฟอร์ม ME-F-01
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. แบบฟอร์มสรุปงานตอนเช้า — ME-F-01")

form_meta = doc.add_table(rows=2, cols=4)
form_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
form_meta.style = 'Table Grid'
add_table_borders(form_meta)
fm_rows = [
    ("วันที่",       "...... / ...... / ..........",  "เวลา",       "......... น."),
    ("ผู้ปฏิบัติงาน","_______________________________","ลงชื่อ",    "_______________________________"),
]
for i, row_data in enumerate(fm_rows):
    for j, txt in enumerate(row_data):
        set_cell_bg(form_meta.rows[i].cells[j], "D6E4F0" if j%2==0 else "FFFFFF")
        cell_text(form_meta.rows[i].cells[j], txt, size=15,
                  bold=(j%2==0), color=DARK_BLUE if j%2==0 else None,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

sum_t = doc.add_table(rows=6, cols=5)
sum_t.alignment = WD_TABLE_ALIGNMENT.CENTER
sum_t.style = 'Table Grid'
add_table_borders(sum_t)
sum_h = ["ระดับ","จำนวน (รายการ)","เสร็จแล้ว","ค้าง","หมายเหตุ"]
for j, h in enumerate(sum_h):
    set_cell_bg(sum_t.rows[0].cells[j], "003366")
    cell_text(sum_t.rows[0].cells[j], h, size=14, bold=True, color=WHITE)
sum_rows = [
    ("D1  วิกฤต",    "FADBD8", RED),
    ("D2  เร่งด่วน", "FEF5E7", ORANGE),
    ("D3  ปกติ",     "EAFAF1", GREEN),
    ("PM",           "F2F3F4", (100,100,100)),
    ("รวมทั้งหมด",   "D6E4F0", DARK_BLUE),
]
for i, (label, bg, col) in enumerate(sum_rows):
    set_cell_bg(sum_t.rows[i+1].cells[0], bg)
    cell_text(sum_t.rows[i+1].cells[0], label, size=14, bold=True, color=col)
    for j in range(1,5):
        set_cell_bg(sum_t.rows[i+1].cells[j], bg)
        cell_text(sum_t.rows[i+1].cells[j], "", size=14)

# ══════════════════════════════════════════════════════════════════════════════
# WARNING BOX
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
warn_t = doc.add_table(rows=1, cols=1)
warn_t.alignment = WD_TABLE_ALIGNMENT.CENTER
warn_t.style = 'Table Grid'
set_cell_bg(warn_t.rows[0].cells[0], "FDEDEC")
cell = warn_t.rows[0].cells[0]
cell.paragraphs[0].clear()
warns = [
    "⚠  ห้ามใช้เครื่องมือแพทย์ที่ยังซ่อมไม่เสร็จกับผู้ป่วย จนกว่าจะผ่านการทดสอบและติดป้าย 'พร้อมใช้งาน' แล้วเท่านั้น",
    "⚠  งาน D1 ทุกรายการต้องรายงานหัวหน้าหน่วยก่อนและหลังดำเนินการ",
    "⚠  บันทึกทุกขั้นตอนการซ่อมลงในระบบหรือแบบฟอร์มทุกครั้ง",
]
for w in warns:
    p = cell.add_paragraph()
    run = p.add_run(w)
    run.font.name = FONT_THAI
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*RED)
    rpr = run._r.get_or_add_rPr()
    cs = OxmlElement('w:cs'); cs.set(qn('w:val'), FONT_THAI)
    rPrCs = rpr.find(qn('w:cs'))
    if rPrCs is not None: rpr.remove(rPrCs)
    rpr.append(cs)
    p.paragraph_format.space_after = Pt(3)

# ── Footer paragraph ─────────────────────────────────────────────────────────
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot.add_run("จัดทำโดย: หน่วยวิศวกรรมการแพทย์  |  อ้างอิง: มาตรฐาน HA, ISO 13485, กระทรวงสาธารณสุข  |  ทบทวนทุก 1 ปี")
run.font.name  = FONT_THAI
run.font.size  = Pt(13)
run.font.italic = True
run.font.color.rgb = RGBColor(127,127,127)
rpr = run._r.get_or_add_rPr()
cs = OxmlElement('w:cs'); cs.set(qn('w:val'), FONT_THAI)
rPrCs = rpr.find(qn('w:cs'))
if rPrCs is not None: rpr.remove(rPrCs)
rpr.append(cs)
foot.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
foot.paragraph_format.line_spacing = 1.5

doc.save(str(WORD_OUT))
print(f"Word saved: {WORD_OUT}")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — สร้าง PDF (enhanced)
# ══════════════════════════════════════════════════════════════════════════════
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Image as RLImage, KeepTogether
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

for rname, rfile in [("Sarabun","Sarabun-Regular"),("Sarabun-Bold","Sarabun-Bold"),("Sarabun-Italic","Sarabun-Italic")]:
    pdfmetrics.registerFont(TTFont(rname, str(FONTS/f"{rfile}.ttf")))

W_PDF, H_PDF = A4
LEFT = RIGHT = 20*mm
TOP  = BOTTOM = 20*mm
CW   = W_PDF - LEFT - RIGHT   # content width

# HA/ISO standard sizes: body=16pt Thai = ~11.3pt RL (1pt=1.333px; but RL uses points directly)
# We map: body→11pt, h1→14pt, h2→12pt (RL points ≈ Thai 16/18/20pt display)
def S(name="body", size=11, bold=False, italic=False, color=colors.black,
       align=TA_LEFT, leading_mult=1.5, sb=0, sa=2):
    fn = "Sarabun-Bold" if bold else ("Sarabun-Italic" if italic else "Sarabun")
    return ParagraphStyle(name, fontName=fn, fontSize=size, textColor=color,
                          alignment=align, leading=size*leading_mult,
                          spaceBefore=sb*mm, spaceAfter=sa*mm)

s_body  = S(size=11)
s_small = S(size=9, color=colors.HexColor("#7F8C8D"))
s_bold  = S(size=11, bold=True)
s_h1    = S(size=14, bold=True, color=colors.HexColor("#003366"), sb=4, sa=2)
s_h2    = S(size=12, bold=True, color=colors.HexColor("#005A9C"), sb=3, sa=1)
s_th    = S(size=9,  bold=True, color=colors.white, align=TA_CENTER)
s_td    = S(size=9,  align=TA_CENTER)
s_td_l  = S(size=9,  align=TA_LEFT)
s_warn  = S(size=9,  bold=True, color=colors.HexColor("#C0392B"))
s_foot  = S(size=8,  italic=True, color=colors.HexColor("#7F8C8D"), align=TA_CENTER)

DARK_B = colors.HexColor("#003366")
MED_B  = colors.HexColor("#005A9C")
LT_B   = colors.HexColor("#D6E4F0")
RED_C  = colors.HexColor("#C0392B")
ORA_C  = colors.HexColor("#E67E22")
GRN_C  = colors.HexColor("#1E8449")
GRY_C  = colors.HexColor("#F2F3F4")
WHITE_C= colors.white

def tbl_style(header_bg=DARK_B, row_bgs=None):
    base = [
        ("BACKGROUND", (0,0), (-1,0), header_bg),
        ("TEXTCOLOR",  (0,0), (-1,0), WHITE_C),
        ("GRID",       (0,0), (-1,-1), 0.4, colors.HexColor("#BDBDBD")),
        ("TOPPADDING",    (0,0),(-1,-1), 4),
        ("BOTTOMPADDING", (0,0),(-1,-1), 4),
        ("LEFTPADDING",   (0,0),(-1,-1), 5),
        ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
    ]
    if row_bgs:
        for i, bg in enumerate(row_bgs):
            base.append(("BACKGROUND",(0,i+1),(-1,i+1), bg))
    return TableStyle(base)

def HR():
    return HRFlowable(width="100%", thickness=2, color=DARK_B, spaceAfter=4)

story = []

# BANNER
story.append(RLImage(str(banner_path), width=CW, height=CW*200/1754))
story.append(Spacer(1, 4*mm))

# META
meta_d = [
    [Paragraph("ฉบับที่",s_th), Paragraph("001",s_td), Paragraph("วันที่บังคับใช้",s_th), Paragraph("มิถุนายน 2569",s_td)],
    [Paragraph("แผนก",s_th),    Paragraph("วิศวกรรมการแพทย์ / ซ่อมบำรุงเครื่องมือแพทย์",s_td_l), Paragraph("เวลาปฏิบัติงาน",s_th), Paragraph("08:30 – 16:30 น.",s_td)],
    [Paragraph("วันทำการ",s_th),Paragraph("จันทร์ – ศุกร์ (วันราชการ)",s_td_l), Paragraph("ผู้รับผิดชอบ",s_th), Paragraph("ช่างเครื่องมือแพทย์ /\nวิศวกรชีวการแพทย์",s_td_l)],
]
mt = Table(meta_d, colWidths=[28*mm, 65*mm, 34*mm, 43*mm])
mt.setStyle(tbl_style(DARK_B))
mt.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,-1),colors.white),
    ("BACKGROUND",(0,0),(0,-1),LT_B),
    ("BACKGROUND",(2,0),(2,-1),LT_B),
    ("FONTNAME",(0,0),(0,-1),"Sarabun-Bold"),
    ("FONTNAME",(2,0),(2,-1),"Sarabun-Bold"),
    ("FONTSIZE",(0,0),(-1,-1),9),
    ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [mt, Spacer(1,4*mm)]

# 1. วัตถุประสงค์
story += [HR(), Paragraph("1. วัตถุประสงค์", s_h1),
    Paragraph("เพื่อคัดกรองและจัดลำดับความสำคัญของงานซ่อมเครื่องมือแพทย์ก่อนเริ่มปฏิบัติงานทุกวัน ให้งานที่เร่งด่วนและส่งผลต่อความปลอดภัยของผู้ป่วยได้รับการแก้ไขก่อนเป็นลำดับแรก ตามมาตรฐาน HA และ ISO 13485", s_body)]

# 2. Timeline
story += [Spacer(1,3*mm), HR(), Paragraph("2. ตารางปฏิบัติงานประจำวัน", s_h1),
    RLImage(str(timeline_path), width=CW, height=CW*160/1600),
    Spacer(1,3*mm)]

# 3. ระดับเร่งด่วน
story += [HR(), Paragraph("3. เกณฑ์จำแนกระดับความเร่งด่วน", s_h1)]
urg_data_p = [
    [Paragraph("ระดับ",s_th),Paragraph("รหัส",s_th),Paragraph("ความหมาย",s_th),Paragraph("ระยะเวลา",s_th),Paragraph("ช่วงเวลา",s_th)],
    [Paragraph("วิกฤต",S(9,bold=True,color=WHITE_C,align=TA_CENTER)),Paragraph("D1",S(9,bold=True,color=WHITE_C,align=TA_CENTER)),
     Paragraph("เครื่องมือเกี่ยวข้องกับชีวิตผู้ป่วยโดยตรง หรือหน่วย ICU/ER/OR",s_td_l),
     Paragraph("ภายใน 1 ชั่วโมง",S(9,bold=True,color=WHITE_C,align=TA_CENTER)),Paragraph("09:00–10:00",S(9,bold=True,color=WHITE_C,align=TA_CENTER))],
    [Paragraph("เร่งด่วน",s_td),Paragraph("D2",s_td),Paragraph("กระทบการรักษา/วินิจฉัย แต่มีเครื่องสำรอง",s_td_l),Paragraph("ภายใน 4 ชั่วโมง",s_td),Paragraph("ก่อน 12:00",s_td)],
    [Paragraph("ปกติ",s_td),Paragraph("D3",s_td),Paragraph("ไม่กระทบการรักษาในทันที",s_td_l),Paragraph("ภายใน 24–48 ชม.",s_td),Paragraph("ตามลำดับ",s_td)],
    [Paragraph("แผน PM",s_td),Paragraph("PM",s_td),Paragraph("งาน Preventive Maintenance ตามกำหนดการ",s_td_l),Paragraph("ตามแผนที่กำหนด",s_td),Paragraph("ตามแผน",s_td)],
]
ut = Table(urg_data_p, colWidths=[20*mm,13*mm,72*mm,30*mm,CW-135*mm])
ut.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,0),DARK_B),
    ("BACKGROUND",(0,1),(-1,1),RED_C),
    ("BACKGROUND",(0,2),(-1,2),colors.HexColor("#FEF5E7")),
    ("BACKGROUND",(0,3),(-1,3),colors.HexColor("#EAFAF1")),
    ("BACKGROUND",(0,4),(-1,4),GRY_C),
    ("TEXTCOLOR",(0,1),(-1,1),WHITE_C),
    ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [ut, Spacer(1,3*mm)]

# 4. ขั้นตอน
story += [HR(), Paragraph("4. ขั้นตอนการตรวจเช็คตอนเช้า — Morning Round (08:30 – 09:00 น.)", s_h1)]
for step_no, title, bullets in steps:
    hdr = [[Paragraph(step_no, S(9,bold=True,color=WHITE_C,align=TA_CENTER)),
            Paragraph(title,   S(10,bold=True,color=DARK_B))]]
    ht = Table(hdr, colWidths=[16*mm, CW-16*mm])
    ht.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,0),MED_B),("BACKGROUND",(1,0),(1,0),LT_B),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("BOX",(0,0),(-1,-1),0.5,MED_B),
    ]))
    bl = [ht] + [Paragraph(f"   ✓  {b}", s_body) for b in bullets] + [Spacer(1,1*mm)]
    story.append(KeepTogether(bl))

# 5. คำถามคัดกรอง
story += [Spacer(1,2*mm), HR(), Paragraph("5. คำถามคัดกรองความเร่งด่วน (ถามผู้แจ้งซ่อม)", s_h1)]
q_data_p = [[Paragraph("ข้อ",s_th), Paragraph("คำถาม",s_th), Paragraph("ตอบ ใช่ → ระดับ",s_th)]]
for no,q,lv in [("1","เครื่องมือนี้ใช้กับผู้ป่วยโดยตรงขณะนี้หรือไม่?","D1"),
                ("2","อยู่ในหน่วย ICU / ER / OR / NICU หรือไม่?","D1"),
                ("3","มีผู้ป่วยรอใช้และไม่มีเครื่องสำรองหรือไม่?","D1"),
                ("4","กระทบการวินิจฉัย/รักษา แต่มีเครื่องสำรองหรือไม่?","D2")]:
    c = RED_C if lv=="D1" else ORA_C
    q_data_p.append([Paragraph(no,s_td), Paragraph(q,s_td_l),
                     Paragraph(lv, S(9,bold=True,color=c,align=TA_CENTER))])
qt = Table(q_data_p, colWidths=[10*mm, CW-50*mm, 40*mm])
qt.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,0),DARK_B),
    ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#FDEDEC"),colors.HexColor("#FEF5E7"),colors.HexColor("#FDEDEC"),colors.HexColor("#FEF5E7")]),
    ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [qt, Spacer(1,3*mm)]

# 6. เครื่องมือเสี่ยงสูง
story += [HR(), Paragraph("6. เครื่องมือแพทย์กลุ่มเสี่ยงสูง (แก้ไขก่อนเสมอ)", s_h1)]
hi_d = [[Paragraph(f"⚕  {a}", S(9,color=colors.HexColor("#641E16"))),
         Paragraph(f"⚕  {b}", S(9,color=colors.HexColor("#641E16")))] for a,b in hi]
hit = Table(hi_d, colWidths=[CW/2, CW/2])
hit.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#FDEDEC")),
    ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#F1948A")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),8),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [hit, Spacer(1,3*mm)]

# 7. Escalation
story += [HR(), Paragraph("7. เกณฑ์การยกระดับความเร่งด่วน (Escalation)", s_h1)]
esc_d = [[Paragraph("สถานการณ์",s_th), Paragraph("การดำเนินการ",s_th)]] + [
    [Paragraph(s, s_td_l), Paragraph(a, S(9,bold=True,color=RED_C))]
    for s,a in esc_data]
et = Table(esc_d, colWidths=[80*mm, CW-80*mm])
et.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,0),DARK_B),
    ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,GRY_C]),
    ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [et, Spacer(1,3*mm)]

# 8. แบบฟอร์ม
story += [HR(), Paragraph("8. แบบฟอร์มสรุปงานตอนเช้า — ME-F-01", s_h1)]
fm_d = [
    [Paragraph("วันที่",s_th), Paragraph("...... / ...... / ..........",s_td), Paragraph("เวลา",s_th), Paragraph("......... น.",s_td)],
    [Paragraph("ผู้ปฏิบัติงาน",s_th), Paragraph("_______________________________",s_td), Paragraph("ลงชื่อ",s_th), Paragraph("_______________________________",s_td)],
]
fmt = Table(fm_d, colWidths=[26*mm, 62*mm, 20*mm, 62*mm])
fmt.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(0,-1),LT_B),("BACKGROUND",(2,0),(2,-1),LT_B),
    ("BACKGROUND",(1,0),(1,-1),colors.white),("BACKGROUND",(3,0),(3,-1),colors.white),
    ("FONTNAME",(0,0),(0,-1),"Sarabun-Bold"),("FONTNAME",(2,0),(2,-1),"Sarabun-Bold"),
    ("FONTSIZE",(0,0),(-1,-1),9),
    ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [fmt, Spacer(1,2*mm)]

sum_d = [[Paragraph(h,s_th) for h in ["ระดับ","จำนวน (รายการ)","เสร็จแล้ว","ค้าง","หมายเหตุ"]]]
sbgs = [colors.HexColor("#FADBD8"),colors.HexColor("#FEF5E7"),colors.HexColor("#EAFAF1"),GRY_C,LT_B]
scols = [RED_C,ORA_C,GRN_C,colors.HexColor("#7F8C8D"),DARK_B]
for (lb,_,sc),bg in zip(sum_rows,sbgs):
    sum_d.append([Paragraph(lb,S(9,bold=True,color=sc,align=TA_CENTER)),"","","",""])
sst = Table(sum_d, colWidths=[28*mm,35*mm,28*mm,22*mm,CW-113*mm])
sst.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,0),DARK_B),
    *[("BACKGROUND",(0,i+1),(-1,i+1),bg) for i,bg in enumerate(sbgs)],
    ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
    ("ALIGN",(0,0),(-1,-1),"CENTER"),
]))
story += [sst, Spacer(1,4*mm)]

# WARNING
warn_d = [[Paragraph(
    "⚠  ห้ามใช้เครื่องมือแพทย์ที่ยังซ่อมไม่เสร็จกับผู้ป่วย จนกว่าจะผ่านการทดสอบและติดป้าย 'พร้อมใช้งาน' แล้วเท่านั้น  "
    "|  งาน D1 ทุกรายการต้องรายงานหัวหน้าหน่วยก่อนและหลังดำเนินการ  "
    "|  บันทึกทุกขั้นตอนการซ่อมลงในระบบหรือแบบฟอร์มทุกครั้ง",
    S(8,bold=True,color=RED_C,align=TA_CENTER))]]
wt = Table(warn_d, colWidths=[CW])
wt.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#FDEDEC")),
    ("BOX",(0,0),(-1,-1),1,RED_C),
    ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
    ("LEFTPADDING",(0,0),(-1,-1),8),
]))
story += [wt, Spacer(1,4*mm)]

# FOOTER
foot_d = [[
    Paragraph("จัดทำโดย: หน่วยวิศวกรรมการแพทย์", s_foot),
    Paragraph("อ้างอิง: มาตรฐาน HA | ISO 13485 | กระทรวงสาธารณสุข", S(8,italic=True,color=colors.HexColor("#7F8C8D"),align=TA_CENTER)),
    Paragraph("ทบทวนทุก 1 ปี", S(8,italic=True,color=colors.HexColor("#7F8C8D"),align=TA_RIGHT)),
]]
fot = Table(foot_d, colWidths=[CW/3, CW/3, CW/3])
fot.setStyle(TableStyle([
    ("LINEABOVE",(0,0),(-1,0),0.5,colors.HexColor("#7F8C8D")),
    ("TOPPADDING",(0,0),(-1,-1),4),
]))
story.append(fot)

pdf_doc = SimpleDocTemplate(str(PDF_OUT), pagesize=A4,
                             leftMargin=LEFT, rightMargin=RIGHT,
                             topMargin=TOP,  bottomMargin=BOTTOM)
pdf_doc.build(story)
print(f"PDF saved: {PDF_OUT}")
print("\nเสร็จสมบูรณ์ทั้ง 2 ไฟล์!")
