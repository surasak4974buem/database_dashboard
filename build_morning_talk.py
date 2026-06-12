#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
โครงการ Morning Talk — ฝ่ายวิศวกรรมการแพทย์ โรงพยาบาลนครพิงค์
สร้าง Word + PDF พร้อม auto-versioning และ auto-push GitHub
"""

import os, subprocess, urllib.request
from datetime import datetime
from pathlib import Path

BASE   = Path("/home/user/database_dashboard")
FONTS  = Path("/tmp/fonts")
FONTS.mkdir(parents=True, exist_ok=True)
IMGS   = BASE / "assets"
IMGS.mkdir(parents=True, exist_ok=True)

BASENAME = "โครงการ_Morning_Talk_นครพิงค์"

FONT_URLS = {
    "Sarabun-Regular": "https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Regular.ttf",
    "Sarabun-Bold":    "https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Bold.ttf",
    "Sarabun-Italic":  "https://github.com/google/fonts/raw/main/ofl/sarabun/Sarabun-Italic.ttf",
}
for name, url in FONT_URLS.items():
    path = FONTS / f"{name}.ttf"
    if not path.exists():
        print(f"  downloading {name}...")
        urllib.request.urlretrieve(url, path)

def versioned(base: Path, suffix: str) -> Path:
    c = base.parent / f"{base.name}{suffix}"
    if not c.exists(): return c
    v = 2
    while True:
        c = base.parent / f"{base.name}_v{v}{suffix}"
        if not c.exists(): return c
        v += 1

WORD_OUT = versioned(BASE / BASENAME, ".docx")
PDF_OUT  = versioned(BASE / BASENAME, ".pdf")
print(f"Word → {WORD_OUT.name}")
print(f"PDF  → {PDF_OUT.name}")

# ── สี ──────────────────────────────────────────────────────────────────────
DARK_BLUE  = (0,   51, 102)
MED_BLUE   = (0,   90, 156)
LIGHT_BLUE = (214, 228, 240)
RED        = (192,  57,  43)
ORANGE     = (230, 126,  34)
GREEN      = ( 30, 132,  73)
TEAL       = ( 23, 165, 137)
PURPLE     = (118,  68, 138)
WHITE      = (255, 255, 255)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Banner image
# ══════════════════════════════════════════════════════════════════════════════
from PIL import Image, ImageDraw, ImageFont

def make_mt_banner():
    W, H = 1754, 220
    img = Image.new("RGB", (W, H), DARK_BLUE)
    d   = ImageDraw.Draw(img)
    for i in range(W):
        r = int(DARK_BLUE[0] + (MED_BLUE[0]-DARK_BLUE[0])*i/W)
        g = int(DARK_BLUE[1] + (MED_BLUE[1]-DARK_BLUE[1])*i/W)
        b = int(DARK_BLUE[2] + (MED_BLUE[2]-DARK_BLUE[2])*i/W)
        d.line([(i,0),(i,H)], fill=(r,g,b))
    d.rectangle([0, H-14, W, H], fill=TEAL)
    d.rectangle([0, 0, W, 6], fill=ORANGE)
    try:
        f1 = ImageFont.truetype(str(FONTS/"Sarabun-Bold.ttf"),    80)
        f2 = ImageFont.truetype(str(FONTS/"Sarabun-Regular.ttf"), 42)
        f3 = ImageFont.truetype(str(FONTS/"Sarabun-Regular.ttf"), 34)
    except: f1=f2=f3=ImageFont.load_default()
    lines = [
        (f1, "โครงการประชุมเช้าย่อย (Morning Talk)", WHITE, 18),
        (f2, "หน่วยซ่อมบำรุงเครื่องมือแพทย์  |  ฝ่ายวิศวกรรมการแพทย์", (214,228,240), 108),
        (f3, "โรงพยาบาลนครพิงค์  จังหวัดเชียงใหม่", (180,210,240), 158),
    ]
    for fnt, txt, col, y in lines:
        bb = d.textbbox((0,0), txt, font=fnt)
        d.text(((W-(bb[2]-bb[0]))//2, y), txt, font=fnt, fill=col)
    p = IMGS / "mt_banner.png"
    img.save(p); return p

def make_flow_chart():
    """ผังขั้นตอน Morning Talk 10 นาที"""
    W, H = 1600, 260
    img = Image.new("RGB", (W, H), (250,250,250))
    d   = ImageDraw.Draw(img)
    try:
        fb = ImageFont.truetype(str(FONTS/"Sarabun-Bold.ttf"),    28)
        fn = ImageFont.truetype(str(FONTS/"Sarabun-Regular.ttf"), 24)
        ft = ImageFont.truetype(str(FONTS/"Sarabun-Bold.ttf"),    22)
    except: fb=fn=ft=ImageFont.load_default()

    steps = [
        (MED_BLUE,  "เปิดประชุม\nผู้นำวันนี้",    "1 นาที"),
        (TEAL,      "งานค้าง\nเวรบ่าย/ดึก",      "2 นาที"),
        (ORANGE,    "แผนงานวันนี้\nPM/CAL/ซ่อม",  "3 นาที"),
        (PURPLE,    "ปัญหา &\nข้อเสนอแนะ",        "2 นาที"),
        (GREEN,     "ทบทวน\nเป้าหมาย",            "1 นาที"),
        (RED,       "มอบหมายงาน\nปิดประชุม",       "1 นาที"),
    ]
    box_w = 220; gap = 40
    total_w = len(steps)*box_w + (len(steps)-1)*gap
    start_x = (W - total_w) // 2
    y1, y2 = 60, 200

    title = "ขั้นตอน Morning Talk (รวม ≤ 10 นาที)"
    bb = d.textbbox((0,0), title, font=fb)
    d.text(((W-(bb[2]-bb[0]))//2, 10), title, font=fb, fill=DARK_BLUE)

    for i, (col, label, dur) in enumerate(steps):
        x1 = start_x + i*(box_w+gap)
        x2 = x1 + box_w
        # box
        d.rounded_rectangle([x1, y1, x2, y2], radius=12, fill=col)
        # arrow
        if i < len(steps)-1:
            ax = x2+2; ay = (y1+y2)//2
            d.polygon([(ax,ay-10),(ax+gap-4,ay),(ax,ay+10)], fill=(180,180,180))
        # label
        for j, ln in enumerate(label.split("\n")):
            bb = d.textbbox((0,0), ln, font=fn)
            tx = x1 + (box_w-(bb[2]-bb[0]))//2
            ty = y1 + 18 + j*32
            d.text((tx, ty), ln, font=fn, fill=WHITE)
        # duration badge
        bb = d.textbbox((0,0), dur, font=ft)
        bw = bb[2]-bb[0]+12
        bx1 = x1+(box_w-bw)//2; bx2 = bx1+bw
        d.rounded_rectangle([bx1, y2-28, bx2, y2-6], radius=8, fill=(255,255,255,180))
        d.text((bx1+6, y2-27), dur, font=ft, fill=col)

    p = IMGS / "mt_flowchart.png"
    img.save(p); return p

def make_shift_diagram():
    """แผนภาพเงื่อนไขเวรศูนย์ยืม"""
    W, H = 1400, 300
    img = Image.new("RGB", (W, H), (250,250,250))
    d   = ImageDraw.Draw(img)
    try:
        fb = ImageFont.truetype(str(FONTS/"Sarabun-Bold.ttf"),    30)
        fn = ImageFont.truetype(str(FONTS/"Sarabun-Regular.ttf"), 26)
    except: fb=fn=ImageFont.load_default()

    title = "เงื่อนไขการเข้าร่วม Morning Talk ตามตารางเวรศูนย์ยืม"
    bb = d.textbbox((0,0), title, font=fb)
    d.text(((W-(bb[2]-bb[0]))//2, 10), title, font=fb, fill=DARK_BLUE)

    cases = [
        (GREEN,  "กรณี A", "ช่างไม่ได้ขึ้นเวรดึก\nเมื่อคืน", "เข้าร่วม\nMorning Talk ✓", 120),
        (RED,    "กรณี B", "ช่างขึ้นเวรดึก\nศูนย์ยืมเมื่อคืน", "ไม่ต้องเข้าร่วม\nใช้ LINE แทน", 120),
        (ORANGE, "กรณี C", "ช่างลาป่วย/\nลาพักร้อน", "แจ้งหัวหน้า\nล่วงหน้า", 120),
    ]
    bw = 360; gap = 50
    total_w = len(cases)*bw+(len(cases)-1)*gap
    sx = (W-total_w)//2
    for i, (col, tag, cond, result, h) in enumerate(cases):
        x1 = sx + i*(bw+gap); x2 = x1+bw
        # condition box
        d.rounded_rectangle([x1, 55, x2, 55+h], radius=14, fill=col)
        bb = d.textbbox((0,0), tag, font=fb)
        d.text((x1+(bw-(bb[2]-bb[0]))//2, 62), tag, font=fb, fill=WHITE)
        for j, ln in enumerate(cond.split("\n")):
            bb = d.textbbox((0,0), ln, font=fn)
            d.text((x1+(bw-(bb[2]-bb[0]))//2, 102+j*32), ln, font=fn, fill=WHITE)
        # arrow
        ax = x1+bw//2; d.polygon([(ax-14,195),(ax+14,195),(ax,215)], fill=(120,120,120))
        # result box
        d.rounded_rectangle([x1, 218, x2, 288], radius=10, fill=LIGHT_BLUE)
        for j, ln in enumerate(result.split("\n")):
            bb = d.textbbox((0,0), ln, font=fn)
            rc = GREEN if "✓" in ln else (RED if "LINE" in ln else DARK_BLUE)
            d.text((x1+(bw-(bb[2]-bb[0]))//2, 225+j*30), ln, font=fn, fill=rc)

    p = IMGS / "mt_shift_diagram.png"
    img.save(p); return p

print("สร้างรูปภาพ...")
banner_path  = make_mt_banner()
flow_path    = make_flow_chart()
shift_path   = make_shift_diagram()
print("  ✓ รูปภาพเสร็จ")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — Word Document
# ══════════════════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
sec = doc.sections[0]
sec.page_width    = Cm(21.0); sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(2.5);  sec.right_margin  = Cm(2.0)
sec.top_margin    = Cm(2.5);  sec.bottom_margin = Cm(2.0)

FONT_THAI = "TH Sarabun New"

def cs_fix(run):
    rpr = run._r.get_or_add_rPr()
    cs  = OxmlElement('w:cs'); cs.set(qn('w:val'), FONT_THAI)
    old = rpr.find(qn('w:cs'))
    if old is not None: rpr.remove(old)
    rpr.append(cs)

def add_run(para, text, size=16, bold=False, italic=False, color=None):
    run = para.add_run(text)
    run.font.name   = FONT_THAI
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)
    cs_fix(run)
    return run

def new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4, ls=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before       = Pt(sb)
    p.paragraph_format.space_after        = Pt(sa)
    p.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing       = ls
    return p

def heading(doc, text, level=1):
    cfg = {1:(20,DARK_BLUE,8,4), 2:(18,MED_BLUE,6,3), 3:(16,DARK_BLUE,4,2)}
    sz, col, sb, sa = cfg.get(level,(16,DARK_BLUE,4,2))
    p = new_para(doc, sb=sb, sa=sa)
    add_run(p, text, size=sz, bold=True, color=col)
    if level == 1:
        ppr  = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'003366')
        pbdr.append(bot); ppr.append(pbdr)
    return p

def body(doc, text, indent=0.0, size=16, sa=4):
    p = new_para(doc, sa=sa)
    p.paragraph_format.left_indent = Cm(indent)
    add_run(p, text, size=size)
    return p

def bullet(doc, text, size=16):
    p = new_para(doc, sa=3)
    p.paragraph_format.left_indent   = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    add_run(p, "◆  ", size=size, color=MED_BLUE)
    add_run(p, text, size=size)
    return p

def sub_bullet(doc, text, size=15):
    p = new_para(doc, sa=2)
    p.paragraph_format.left_indent   = Cm(1.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    add_run(p, "–  ", size=size, color=TEAL)
    add_run(p, text, size=size)

def set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
    tcPr.append(shd)

def cell_txt(cell, text, size=14, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = FONT_THAI; run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    cs_fix(run)

def tbl_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tb = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),'BDBDBD')
        tb.append(el)
    tblPr.append(tb)

def add_img(doc, path, w_cm=16.5):
    doc.add_picture(str(path), width=Cm(w_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── ปก / Banner ─────────────────────────────────────────────────────────────
add_img(doc, banner_path)
doc.add_paragraph()

# ── ข้อมูลโครงการ ─────────────────────────────────────────────────────────
info = doc.add_table(rows=5, cols=4)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info.style = 'Table Grid'; tbl_borders(info)
info_rows = [
    ("ชื่อโครงการ",   "โครงการประชุมเช้าย่อย (Morning Talk)",           "ปีงบประมาณ",  "2568–2569"),
    ("หน่วยงาน",      "ฝ่ายวิศวกรรมการแพทย์ โรงพยาบาลนครพิงค์",        "ระยะเวลา",    "ต.ค. 2568 – ก.ย. 2569"),
    ("ผู้รับผิดชอบ",  "หัวหน้าหน่วยซ่อมบำรุงเครื่องมือแพทย์",           "วันดำเนินการ","จันทร์ – ศุกร์ (วันราชการ)"),
    ("เวลา",          "08:30 – 08:40 น. (ไม่เกิน 10 นาที)",             "สถานที่",     "ห้องปฏิบัติการวิศวกรรมการแพทย์"),
    ("กลุ่มเป้าหมาย","ช่างเครื่องมือแพทย์ / วิศวกรชีวการแพทย์ทุกคน",   "จำนวน",       "ทุกคนในหน่วย"),
]
for i, (k1,v1,k2,v2) in enumerate(info_rows):
    cells = info.rows[i].cells
    set_bg(cells[0],"D6E4F0"); set_bg(cells[2],"D6E4F0")
    set_bg(cells[1],"FFFFFF"); set_bg(cells[3],"FFFFFF")
    cell_txt(cells[0],k1,14,True,DARK_BLUE); cell_txt(cells[1],v1,14,align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_txt(cells[2],k2,14,True,DARK_BLUE); cell_txt(cells[3],v2,14,align=WD_ALIGN_PARAGRAPH.LEFT)
    for j,w in enumerate([Cm(3.0),Cm(6.5),Cm(3.0),Cm(4.0)]): cells[j].width=w

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"1. หลักการและเหตุผล")
body(doc,"ช่างซ่อมบำรุงเครื่องมือแพทย์ โรงพยาบาลนครพิงค์ ปฏิบัติงานเฉพาะในเวลาราชการ (08:30–16:30 น. วันจันทร์–ศุกร์) ในขณะที่โรงพยาบาลให้บริการตลอด 24 ชั่วโมง เมื่อเครื่องมือแพทย์ชำรุดนอกเวลาราชการ จะใช้ระบบ ศูนย์ยืม-คืนเครื่องมือแพทย์ หมุนเวียนทดแทนชั่วคราว")
doc.add_paragraph()
body(doc,"ปัญหาที่พบในปัจจุบัน ได้แก่:")
for t in ["ข้อมูลงานค้างจากเวรบ่ายหรือเวรดึกไม่ถูกสื่อสารอย่างเป็นระบบสู่ช่างเช้า",
          "การแจ้งงานผ่าน LINE กลุ่ม ไม่สามารถยืนยันได้ว่าผู้รับผิดชอบได้อ่านและรับทราบเมื่อใด",
          "ขาดเวทีพูดคุยแลกเปลี่ยนปัญหา แผน PM แผน CAL และงานซ่อมเร่งด่วนประจำวัน",
          "บุคลากรยังขาดโอกาสฝึกทักษะการนำเสนอและภาวะผู้นำในระดับทีม",
          "เป้าหมายของหน่วยงานไม่ได้รับการทบทวนอย่างสม่ำเสมอ"]: bullet(doc,t)

doc.add_paragraph()
body(doc,"โครงการ Morning Talk จึงเป็นการประชุมย่อยแบบยืน (Stand-up Meeting) ก่อนเริ่มงาน ไม่เกิน 10 นาที เพื่อแก้ปัญหาข้างต้นอย่างเป็นรูปธรรม พร้อมเสริมสร้างวัฒนธรรมการสื่อสารและภาวะผู้นำในทีม")

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"2. วัตถุประสงค์")
objs = [
    "เพื่อสื่อสารข้อมูลงานค้างจากเวรบ่าย/เวรดึกสู่ทีมเช้าอย่างครบถ้วนและทันท่วงที",
    "เพื่อวางแผนและจัดลำดับความสำคัญของงาน PM, CAL, งานซ่อม และงานค้างประจำวัน",
    "เพื่อให้บุคลากรทุกคนได้ฝึกภาวะผู้นำผ่านการหมุนเวียนนำประชุม",
    "เพื่อทบทวนเป้าหมายและตัวชี้วัดของหน่วยงานอย่างสม่ำเสมอ",
    "เพื่อลดข้อผิดพลาดจากการสื่อสารที่ไม่ครบถ้วนในทีมซ่อมบำรุงเครื่องมือแพทย์",
]
for i,t in enumerate(objs,1): bullet(doc,f"{i}. {t}")

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"3. เป้าหมายและตัวชี้วัดความสำเร็จ")
kpis = [
    ("ความสม่ำเสมอในการจัด",  "Morning Talk",                "≥ 80% ของวันทำการ"),
    ("การหมุนเวียนผู้นำ",     "ทุกคนได้นำอย่างน้อย",        "1 ครั้ง/เดือน"),
    ("งานค้างเวร",            "ถูกส่งต่ออย่างครบถ้วน",      "100%"),
    ("ความพึงพอใจทีม",        "แบบประเมินรายไตรมาส",         "≥ 3.5/5.0"),
    ("ระยะเวลาประชุม",        "ไม่เกิน 10 นาที",             "≥ 90% ของครั้งที่จัด"),
]
t = doc.add_table(rows=len(kpis)+1, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER
t.style='Table Grid'; tbl_borders(t)
for j,(h,w) in enumerate(zip(["ตัวชี้วัด","รายละเอียด","เกณฑ์"],
                               [Cm(4.5),Cm(7.5),Cm(4.5)])):
    set_bg(t.rows[0].cells[j],"003366"); t.rows[0].cells[j].width=w
    cell_txt(t.rows[0].cells[j],h,14,True,WHITE)
for i,(k,v,g) in enumerate(kpis):
    bg = "F2F3F4" if i%2==0 else "FFFFFF"
    for j in range(3): set_bg(t.rows[i+1].cells[j],bg)
    cell_txt(t.rows[i+1].cells[0],k,14,bold=True,color=DARK_BLUE,align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_txt(t.rows[i+1].cells[1],v,14,align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_txt(t.rows[i+1].cells[2],g,14,bold=True,color=GREEN)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"4. รูปแบบและขั้นตอน Morning Talk")
body(doc,"4.1  ลักษณะการประชุม: Stand-up Meeting ยืนหน้ากระดาน Whiteboard หรือบอร์ดแสดงแผนงาน")
body(doc,"4.2  ระยะเวลา: ไม่เกิน 10 นาที  |  เวลา: 08:30 – 08:40 น.")
body(doc,"4.3  ผู้นำ: หมุนเวียนทุกคนในทีม (ตามตารางที่กำหนด)")
doc.add_paragraph()
add_img(doc, flow_path, 16.0)
doc.add_paragraph()

steps_detail = [
    ("นาทีที่ 1",    "เปิดประชุม – ผู้นำวันนี้แนะนำตัวและประกาศจำนวนงานรวม",                         MED_BLUE),
    ("นาทีที่ 2–3",  "รายงานงานค้าง – จากเวรบ่าย เวรดึก ศูนย์ยืม และ LINE กลุ่ม",                   TEAL),
    ("นาทีที่ 4–6",  "แผนงานวันนี้ – PM, Calibration (CAL), งานซ่อม D1/D2/D3, และงานนอกแผน",        ORANGE),
    ("นาทีที่ 7–8",  "ปัญหาและข้อเสนอแนะ – ปัญหาที่ต้องการให้หัวหน้าติดตาม หรือต้องการสนับสนุน",    PURPLE),
    ("นาทีที่ 9",    "ทบทวนเป้าหมายหน่วยงาน – ตัวชี้วัดประจำเดือน/ไตรมาส",                          GREEN),
    ("นาทีที่ 10",   "มอบหมายงานและปิดประชุม – ผู้นำสรุปและส่งมอบงาน",                               RED),
]
st = doc.add_table(rows=len(steps_detail)+1, cols=3)
st.alignment=WD_TABLE_ALIGNMENT.CENTER; st.style='Table Grid'; tbl_borders(st)
for j,(h,w) in enumerate(zip(["เวลา","กิจกรรม","หมายเหตุ"],
                              [Cm(2.5),Cm(10.5),Cm(3.5)])):
    set_bg(st.rows[0].cells[j],"003366"); st.rows[0].cells[j].width=w
    cell_txt(st.rows[0].cells[j],h,14,True,WHITE)
for i,(tm,act,col) in enumerate(steps_detail):
    hex_bg = "FFFFFF" if i%2==0 else "F8F9FA"
    set_bg(st.rows[i+1].cells[0],hex_bg); set_bg(st.rows[i+1].cells[1],hex_bg); set_bg(st.rows[i+1].cells[2],hex_bg)
    cell_txt(st.rows[i+1].cells[0],tm,13,bold=True,color=col)
    cell_txt(st.rows[i+1].cells[1],act,13,align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_txt(st.rows[i+1].cells[2],"ผู้นำดำเนินการ" if i in [0,5] else "ทีมร่วมแสดงความเห็น",12,color=(100,100,100))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"5. เงื่อนไขการเข้าร่วมตามตารางเวรศูนย์ยืม")
body(doc,"เพื่อลดแรงกดดันและดูแลสุขภาวะของบุคลากร กำหนดเงื่อนไขดังนี้:")
doc.add_paragraph()
add_img(doc, shift_path, 16.0)
doc.add_paragraph()

cond_data = [
    ("กรณี A","ช่างไม่ได้ขึ้นเวรดึกศูนย์ยืมเมื่อคืน","เข้าร่วม Morning Talk ตามปกติ","003366","EAFAF1","1E8449"),
    ("กรณี B","ช่างขึ้นเวรดึก (00:00–08:00 น.) ศูนย์ยืมเมื่อคืน","ไม่ต้องเข้าร่วม — ส่งข้อมูลผ่าน LINE กลุ่มแทน","C0392B","FDEDEC","C0392B"),
    ("กรณี C","ช่างลาป่วย / ลาพักร้อน / ลากิจ","แจ้งหัวหน้าล่วงหน้า มอบหมายงานผ่าน LINE","E67E22","FEF5E7","E67E22"),
]
ct = doc.add_table(rows=len(cond_data)+1,cols=3)
ct.alignment=WD_TABLE_ALIGNMENT.CENTER; ct.style='Table Grid'; tbl_borders(ct)
for j,(h,w) in enumerate(zip(["กรณี","เงื่อนไข","การปฏิบัติ"],
                               [Cm(2.0),Cm(7.5),Cm(7.0)])):
    set_bg(ct.rows[0].cells[j],"003366"); ct.rows[0].cells[j].width=w
    cell_txt(ct.rows[0].cells[j],h,14,True,WHITE)
for i,(tag,cond,action,tc,bg,ac) in enumerate(cond_data):
    for j in range(3): set_bg(ct.rows[i+1].cells[j],bg)
    r,g,b = tuple(int(tc[k:k+2],16) for k in (0,2,4))
    cell_txt(ct.rows[i+1].cells[0],tag,13,True,(r,g,b))
    cell_txt(ct.rows[i+1].cells[1],cond,13,align=WD_ALIGN_PARAGRAPH.LEFT)
    r2,g2,b2 = tuple(int(ac[k:k+2],16) for k in (0,2,4))
    cell_txt(ct.rows[i+1].cells[2],action,13,bold=True,color=(r2,g2,b2),align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

body(doc,"หมายเหตุ: ช่วงเริ่มต้นโครงการ (Phase 1) Morning Talk จะดำเนินการเฉพาะวันที่ไม่มีช่างขึ้นเวรดึกศูนย์ยืม เพื่อลดแรงกดดันให้ผู้ปฏิบัติงาน และจะปรับเป็น Phase 2 เมื่อทีมพร้อมและมีความเข้าใจในระบบ",size=14)

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"6. การหมุนเวียนผู้นำ (Rotating Leader)")
body(doc,"6.1  วัตถุประสงค์: ฝึกทักษะการสื่อสาร การนำเสนอ และภาวะผู้นำให้กับทุกคนในทีม")
body(doc,"6.2  รูปแบบตารางหมุนเวียน:")
for t in ["จัดทำตารางผู้นำล่วงหน้า 1 เดือน โดยหัวหน้าหน่วย",
          "สับเปลี่ยนตามตัวอักษรชื่อ หรือตามความสมัครใจ",
          "กรณีผู้นำติดเวรดึกหรือลา ให้ผู้นำถัดไปรับหน้าที่แทน",
          "บันทึกชื่อผู้นำในแบบฟอร์ม MT-F-01 ทุกครั้ง"]:
    sub_bullet(doc,t)

doc.add_paragraph()
body(doc,"6.3  หัวข้อที่ผู้นำอาจนำเสนอเพิ่มเติม (เลือก 1 หัวข้อต่อครั้ง ภายใน 1 นาที):")
extras = [
    "ปัญหาที่พบบ่อยและแนวทางแก้ไข",
    "ความรู้ทางเทคนิคสั้นๆ เกี่ยวกับเครื่องมือแพทย์",
    "ข่าวสารอบรม / สัมมนา / มาตรฐานใหม่",
    "ข้อเสนอแนะเพื่อพัฒนางาน",
    "การทบทวนความปลอดภัยในการปฏิบัติงาน (Safety Moment)",
]
for e in extras: sub_bullet(doc,e)

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"7. การแก้ปัญหา LINE กลุ่ม")
body(doc,"แม้จะมี LINE กลุ่มสำหรับสื่อสาร แต่ไม่สามารถยืนยันได้ว่าผู้รับผิดชอบงานจะอ่านและตอบสนองเมื่อใด Morning Talk จึงทำหน้าที่เป็น 'จุดยืนยันการรับรู้' (Acknowledgement Point) ที่เชื่อถือได้ในทุกวันทำการ")
for t in ["งานที่ส่งใน LINE เวรดึก/บ่าย จะถูกยืนยันรับทราบใน Morning Talk เช้าวันถัดไป",
          "ผู้นำ Morning Talk ทำหน้าที่ตรวจสอบว่าทุกงานใน LINE ได้รับการมอบหมายแล้ว",
          "กรณีเร่งด่วนนอกเวลา ยังคงใช้ LINE + โทรศัพท์ตามระเบียบเดิม"]:
    bullet(doc,t)

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"8. แผนดำเนินโครงการ (Gantt)")
gantt_h = ["กิจกรรม","ต.ค.","พ.ย.","ธ.ค.","ม.ค.","ก.พ.","มี.ค.","เม.ย.–ก.ย."]
gantt_d = [
    ("จัดทำคู่มือและแบบฟอร์ม",       "●","●","","","","",""),
    ("ชี้แจงทีมและทดลองใช้ (Phase 1)","","●","●","","","",""),
    ("ปรับปรุงและประเมินผลครั้งที่ 1","","","●","","","",""),
    ("ขยาย Phase 2 (ทุกวันทำการ)",    "","","","●","●","",""),
    ("ประเมินผลไตรมาส 2",             "","","","","","●",""),
    ("สรุปโครงการและรายงาน",          "","","","","","","●"),
]
gt = doc.add_table(rows=len(gantt_d)+1, cols=len(gantt_h))
gt.alignment=WD_TABLE_ALIGNMENT.CENTER; gt.style='Table Grid'; tbl_borders(gt)
for j,(h,w) in enumerate(zip(gantt_h,[Cm(5.0)]+[Cm(1.3)]*6+[Cm(2.2)])):
    set_bg(gt.rows[0].cells[j],"003366"); gt.rows[0].cells[j].width=w
    cell_txt(gt.rows[0].cells[j],h,12,True,WHITE)
for i,(row) in enumerate(gantt_d):
    bg = "F2F3F4" if i%2==0 else "FFFFFF"
    for j in range(len(gantt_h)): set_bg(gt.rows[i+1].cells[j],bg)
    cell_txt(gt.rows[i+1].cells[0],row[0],12,align=WD_ALIGN_PARAGRAPH.LEFT)
    for j in range(1,len(gantt_h)):
        mark = row[j] if j<len(row) else ""
        cell_txt(gt.rows[i+1].cells[j],mark,14,bold=True,
                 color=TEAL if mark=="●" else (200,200,200))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"9. แบบฟอร์มที่ใช้ร่วมโครงการ")
forms = [
    ("MT-F-01","บันทึก Morning Talk ประจำวัน","กรอกทุกครั้งที่มีการประชุม"),
    ("MT-F-02","ตารางหมุนเวียนผู้นำรายเดือน","จัดทำล่วงหน้า 1 เดือน"),
    ("MT-F-03","แบบประเมินความพึงพอใจ",      "ทบทวนรายไตรมาส"),
    ("ME-F-01","สรุปงานซ่อมประจำวัน",         "ใช้ร่วมกับคู่มือการตรวจเช็คงานซ่อม"),
]
ft2 = doc.add_table(rows=len(forms)+1,cols=3)
ft2.alignment=WD_TABLE_ALIGNMENT.CENTER; ft2.style='Table Grid'; tbl_borders(ft2)
for j,(h,w) in enumerate(zip(["รหัสแบบฟอร์ม","ชื่อแบบฟอร์ม","ความถี่การใช้"],
                               [Cm(2.5),Cm(8.5),Cm(5.5)])):
    set_bg(ft2.rows[0].cells[j],"003366"); ft2.rows[0].cells[j].width=w
    cell_txt(ft2.rows[0].cells[j],h,14,True,WHITE)
for i,(code,name,freq) in enumerate(forms):
    bg = "F2F3F4" if i%2==0 else "FFFFFF"
    for j in range(3): set_bg(ft2.rows[i+1].cells[j],bg)
    cell_txt(ft2.rows[i+1].cells[0],code,13,True,TEAL)
    cell_txt(ft2.rows[i+1].cells[1],name,13,align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_txt(ft2.rows[i+1].cells[2],freq,13)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
heading(doc,"10. ผลที่คาดว่าจะได้รับ")
results = [
    "ข้อมูลงานค้างจากเวรดึก/บ่ายถูกส่งต่ออย่างครบถ้วนและทันเวลา ลดความเสี่ยงต่อความปลอดภัยผู้ป่วย",
    "บุคลากรทุกคนได้รับโอกาสฝึกภาวะผู้นำและทักษะการสื่อสารในที่สาธารณะ",
    "ทีมมีความเข้าใจเป้าหมายของหน่วยงานอย่างสม่ำเสมอ",
    "ลดปัญหาการสื่อสารผ่าน LINE ที่ไม่ได้รับการยืนยัน",
    "สร้างวัฒนธรรมการทำงานเชิงรุก (Proactive) ในทีมซ่อมบำรุงเครื่องมือแพทย์",
    "ลดเวลาซ่อมโดยรวม (TAT) เนื่องจากมีการวางแผนงานที่ดีขึ้น",
]
for r in results: bullet(doc,r)

doc.add_paragraph()

# ── ลงนาม ────────────────────────────────────────────────────────────────────
heading(doc,"11. การอนุมัติโครงการ")
sign_t = doc.add_table(rows=3, cols=3)
sign_t.alignment=WD_TABLE_ALIGNMENT.CENTER; sign_t.style='Table Grid'; tbl_borders(sign_t)
for j,(role,w) in enumerate(zip(["ผู้จัดทำโครงการ","หัวหน้าหน่วยซ่อมบำรุง","ผู้อำนวยการ/หัวหน้าฝ่าย"],
                                  [Cm(5.5),Cm(5.5),Cm(5.5)])):
    set_bg(sign_t.rows[0].cells[j],"D6E4F0"); sign_t.rows[0].cells[j].width=w
    cell_txt(sign_t.rows[0].cells[j],role,13,True,DARK_BLUE)
for row in [1,2]:
    for j in range(3):
        label = ["ชื่อ-สกุล: ..................................",
                 "ลงชื่อ: ................................../\n วันที่: .................................."][row-1]
        cell_txt(sign_t.rows[row].cells[j],label,12)

# ── footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=0)
add_run(fp,"โรงพยาบาลนครพิงค์  |  ฝ่ายวิศวกรรมการแพทย์  |  อ้างอิง: มาตรฐาน HA, ISO 13485, แนวทาง Stand-up Meeting",
        size=13, italic=True, color=(127,127,127))

doc.save(str(WORD_OUT))
print(f"Word saved: {WORD_OUT.name}")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — PDF
# ══════════════════════════════════════════════════════════════════════════════
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, HRFlowable, Image as RLImage, KeepTogether)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

for n,f in [("Sarabun","Sarabun-Regular"),("Sarabun-Bold","Sarabun-Bold"),("Sarabun-Italic","Sarabun-Italic")]:
    pdfmetrics.registerFont(TTFont(n, str(FONTS/f"{f}.ttf")))

WP,HP = A4
LM=RM=20*mm; TM=BM=20*mm; CW=WP-LM-RM

DB  = colors.HexColor("#003366"); MB  = colors.HexColor("#005A9C")
LB  = colors.HexColor("#D6E4F0"); RC  = colors.HexColor("#C0392B")
OC  = colors.HexColor("#E67E22"); GC  = colors.HexColor("#1E8449")
TC  = colors.HexColor("#17A589"); PC  = colors.HexColor("#76448A")
GY  = colors.HexColor("#F2F3F4"); WC  = colors.white

def PS(size=11, bold=False, italic=False, color=colors.black,
       align=TA_LEFT, lm=1.5, sb=0, sa=2):
    fn = "Sarabun-Bold" if bold else ("Sarabun-Italic" if italic else "Sarabun")
    return ParagraphStyle("x", fontName=fn, fontSize=size, textColor=color,
                          alignment=align, leading=size*lm, spaceBefore=sb*mm, spaceAfter=sa*mm)

sb_  = PS(11); sh1 = PS(14,True,color=DB,sb=4,sa=2); sh2 = PS(12,True,color=MB,sb=3,sa=1)
sth  = PS(9,True,color=WC,align=TA_CENTER); std=PS(9,align=TA_CENTER); stdl=PS(9,align=TA_LEFT)
sft  = PS(8,italic=True,color=colors.HexColor("#7F8C8D"),align=TA_CENTER)

def HR(): return HRFlowable(width="100%",thickness=2,color=DB,spaceAfter=4)
def IMG(p,w=CW,ratio=1): return RLImage(str(p),width=w,height=w*ratio)

story=[]

# Banner
story += [IMG(banner_path, CW, 220/1754), Spacer(1,4*mm)]

# Meta
md=[
    [Paragraph("ชื่อโครงการ",sth),Paragraph("โครงการประชุมเช้าย่อย (Morning Talk)",stdl),
     Paragraph("ปีงบประมาณ",sth),Paragraph("2568–2569",std)],
    [Paragraph("หน่วยงาน",sth),Paragraph("ฝ่ายวิศวกรรมการแพทย์ รพ.นครพิงค์",stdl),
     Paragraph("ระยะเวลา",sth),Paragraph("ต.ค. 2568 – ก.ย. 2569",std)],
    [Paragraph("เวลา",sth),Paragraph("08:30 – 08:40 น. (≤ 10 นาที)",stdl),
     Paragraph("วันดำเนินการ",sth),Paragraph("จันทร์ – ศุกร์ (วันราชการ)",std)],
]
mt2=Table(md,colWidths=[28*mm,68*mm,32*mm,42*mm])
mt2.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(0,-1),LB),("BACKGROUND",(2,0),(2,-1),LB),
    ("BACKGROUND",(1,0),(1,-1),WC),("BACKGROUND",(3,0),(3,-1),WC),
    ("FONTNAME",(0,0),(0,-1),"Sarabun-Bold"),("FONTNAME",(2,0),(2,-1),"Sarabun-Bold"),
    ("FONTSIZE",(0,0),(-1,-1),9),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [mt2, Spacer(1,4*mm)]

def section(title, content_list):
    items = [HR(), Paragraph(title, sh1)]
    items.extend(content_list)
    return items

# 1. หลักการ
story += section("1. หลักการและเหตุผล", [
    Paragraph("ช่างซ่อมบำรุงเครื่องมือแพทย์ปฏิบัติงานเฉพาะเวลาราชการ ขณะที่โรงพยาบาลให้บริการ 24 ชั่วโมง เมื่อเครื่องมือชำรุดนอกเวลา ใช้ระบบ ศูนย์ยืม-คืน เครื่องมือแพทย์ทดแทน ปัญหาที่พบ ได้แก่ ข้อมูลงานค้างไม่ถูกส่งต่ออย่างเป็นระบบ LINE กลุ่มไม่รับประกันการรับรู้ และขาดเวทีฝึกภาวะผู้นำ โครงการนี้จึงเป็น Stand-up Meeting ก่อนเริ่มงาน ไม่เกิน 10 นาที เพื่อแก้ปัญหาดังกล่าว",sb_),
    Spacer(1,2*mm),
])

# 2. วัตถุประสงค์
obj_rows = [[Paragraph(f"{i}. {t}",stdl)] for i,t in enumerate([
    "สื่อสารงานค้างจากเวรบ่าย/ดึกสู่ทีมเช้าอย่างครบถ้วน",
    "วางแผน PM, CAL, งานซ่อม D1–D3 ประจำวัน",
    "ฝึกภาวะผู้นำผ่านการหมุนเวียนนำประชุม",
    "ทบทวนเป้าหมายและตัวชี้วัดหน่วยงาน",
    "ลดข้อผิดพลาดจากการสื่อสารไม่ครบถ้วน",
],1)]
ot=Table(obj_rows,colWidths=[CW])
ot.setStyle(TableStyle([
    ("ROWBACKGROUNDS",(0,0),(-1,-1),[LB,WC]),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),8),
]))
story += section("2. วัตถุประสงค์", [ot, Spacer(1,3*mm)])

# 3. ขั้นตอน
story += [HR(), Paragraph("3. รูปแบบและขั้นตอน Morning Talk (≤ 10 นาที)", sh1),
    IMG(flow_path, CW, 260/1600), Spacer(1,3*mm)]
step_rows = [[Paragraph(tm,PS(9,True,color=c,align=TA_CENTER)),Paragraph(act,stdl)]
    for tm,act,c in steps_detail]
step_hdr  = [[Paragraph("เวลา",sth),Paragraph("กิจกรรม",sth)]]
step_all  = step_hdr + step_rows
step_t2   = Table(step_all,colWidths=[24*mm,CW-24*mm])
step_t2.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,0),DB),
    ("ROWBACKGROUNDS",(0,1),(-1,-1),[WC,GY]),
    ("GRID",(0,0),(-1,-1),.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [step_t2, Spacer(1,3*mm)]

# 4. เงื่อนไขเวร
story += [HR(), Paragraph("4. เงื่อนไขการเข้าร่วมตามตารางเวรศูนย์ยืม", sh1),
    IMG(shift_path, CW, 300/1400), Spacer(1,2*mm)]
cond_hdr = [[Paragraph("กรณี",sth),Paragraph("เงื่อนไข",sth),Paragraph("การปฏิบัติ",sth)]]
cond_rows2=[]
for tag,cond,action,tc,bg,ac in cond_data:
    r,g,b2 = tuple(int(tc[k:k+2],16) for k in (0,2,4))
    r2,g2,b3 = tuple(int(ac[k:k+2],16) for k in (0,2,4))
    cond_rows2.append([
        Paragraph(tag, PS(9,True,color=colors.Color(r/255,g/255,b2/255),align=TA_CENTER)),
        Paragraph(cond, stdl),
        Paragraph(action, PS(9,True,color=colors.Color(r2/255,g2/255,b3/255))),
    ])
ct2=Table(cond_hdr+cond_rows2,colWidths=[18*mm,80*mm,CW-98*mm])
ct2.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,0),DB),
    ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#EAFAF1"),colors.HexColor("#FDEDEC"),colors.HexColor("#FEF5E7")]),
    ("GRID",(0,0),(-1,-1),.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += [ct2,
    Paragraph("★ Phase 1: Morning Talk เฉพาะวันที่ไม่มีช่างขึ้นเวรดึกศูนย์ยืม เพื่อลดแรงกดดันผู้ปฏิบัติงาน",
              PS(9,True,color=OC,sb=2))]

# 5. KPI
kpi_hdr = [[Paragraph(h,sth) for h in ["ตัวชี้วัด","รายละเอียด","เกณฑ์"]]]
kpi_rows = [[Paragraph(k,PS(9,True,color=DB)),Paragraph(v,stdl),Paragraph(g,PS(9,True,color=GC,align=TA_CENTER))]
    for k,v,g in kpis]
kt=Table(kpi_hdr+kpi_rows,colWidths=[38*mm,72*mm,CW-110*mm])
kt.setStyle(TableStyle([
    ("BACKGROUND",(0,0),(-1,0),DB),
    ("ROWBACKGROUNDS",(0,1),(-1,-1),[GY,WC]),
    ("GRID",(0,0),(-1,-1),.4,colors.HexColor("#BDBDBD")),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),5),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
]))
story += section("5. เป้าหมายและตัวชี้วัด (KPI)", [kt, Spacer(1,3*mm)])

# 6. ผลที่คาดว่าจะได้รับ
res_rows=[[Paragraph(f"✓  {r}",PS(9,color=DB))] for r in results]
rt=Table(res_rows,colWidths=[CW])
rt.setStyle(TableStyle([
    ("ROWBACKGROUNDS",(0,0),(-1,-1),[LB,WC]),
    ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ("LEFTPADDING",(0,0),(-1,-1),8),
]))
story += section("6. ผลที่คาดว่าจะได้รับ", [rt, Spacer(1,3*mm)])

# footer
foot_row=[[Paragraph("โรงพยาบาลนครพิงค์  |  ฝ่ายวิศวกรรมการแพทย์",sft),
           Paragraph("อ้างอิง: มาตรฐาน HA, ISO 13485, Stand-up Meeting Best Practice",PS(8,italic=True,color=colors.HexColor("#7F8C8D"),align=TA_CENTER)),
           Paragraph("ทบทวนทุกปีงบประมาณ",PS(8,italic=True,color=colors.HexColor("#7F8C8D"),align=TA_RIGHT))]]
fott=Table(foot_row,colWidths=[CW/3,CW/3,CW/3])
fott.setStyle(TableStyle([("LINEABOVE",(0,0),(-1,0),.5,colors.HexColor("#7F8C8D")),("TOPPADDING",(0,0),(-1,-1),4)]))
story.append(fott)

pdf_doc=SimpleDocTemplate(str(PDF_OUT),pagesize=A4,leftMargin=LM,rightMargin=RM,topMargin=TM,bottomMargin=BM)
pdf_doc.build(story)
print(f"PDF  saved: {PDF_OUT.name}")
print("\nเสร็จสมบูรณ์ทั้ง 2 ไฟล์!")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — Auto git commit + push
# ══════════════════════════════════════════════════════════════════════════════
def run(cmd): return subprocess.run(cmd,cwd=str(BASE),capture_output=True,text=True)

print("\nกำลัง push ไป GitHub...")
branch = run(["git","rev-parse","--abbrev-ref","HEAD"]).stdout.strip() or "claude/medical-equipment-checklist-86kzld"
now    = datetime.now().strftime("%Y-%m-%d %H:%M")

run(["git","add", WORD_OUT.name, PDF_OUT.name])
msg=(f"เพิ่มโครงการ Morning Talk: {WORD_OUT.name} และ {PDF_OUT.name}\n\n"
     f"สร้างอัตโนมัติโดย build_morning_talk.py — {now}\n"
     f"โรงพยาบาลนครพิงค์ ฝ่ายวิศวกรรมการแพทย์")
cr = run(["git","commit","-m",msg])

if cr.returncode == 0:
    pushed=False
    import time
    for attempt,delay in enumerate([0,2,4,8,16],1):
        if delay: time.sleep(delay); print(f"  retry {attempt}/5...")
        r=subprocess.run(["git","push","-u","origin",branch],cwd=str(BASE),capture_output=True,text=True)
        if r.returncode==0: pushed=True; break
    print(f"  {'✓ push สำเร็จ → branch: '+branch if pushed else '[error] push ไม่สำเร็จ'}")
else:
    print("  ไม่มีการเปลี่ยนแปลง — ข้าม push")
